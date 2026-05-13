import csv
import io
from dataclasses import dataclass
from pathlib import Path

import PyPDF2
import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


@dataclass
class ParsedDocument:
    text: str
    metadata: dict
    page_map: dict[int, str]


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    page_map: dict[int, str] = {}
    full_text: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        page_map[i] = text
        full_text.append(text)
    return ParsedDocument(
        text="\n\n".join(full_text),
        metadata={"filename": filename, "page_count": len(reader.pages)},
        page_map=page_map,
    )


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedDocument(text=text, metadata={"filename": filename}, page_map={1: text})


def parse_txt(file_bytes: bytes, filename: str) -> ParsedDocument:
    text = file_bytes.decode("utf-8", errors="replace")
    return ParsedDocument(text=text, metadata={"filename": filename}, page_map={1: text})


def parse_csv(file_bytes: bytes, filename: str) -> ParsedDocument:
    reader = csv.DictReader(io.StringIO(file_bytes.decode("utf-8", errors="replace")))
    rows = list(reader)
    text = "\n".join([str(row) for row in rows])
    return ParsedDocument(
        text=text,
        metadata={"filename": filename, "row_count": len(rows)},
        page_map={1: text},
    )


async def parse_url(url: str) -> ParsedDocument:
    async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
        resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return ParsedDocument(
        text=text,
        metadata={"source_url": url, "title": soup.title.string if soup.title else url},
        page_map={1: text},
    )


def dispatch_parser(file_bytes: bytes, filename: str, mime: str) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or "pdf" in mime:
        return parse_pdf(file_bytes, filename)
    if ext in (".docx", ".doc"):
        return parse_docx(file_bytes, filename)
    if ext in (".csv",):
        return parse_csv(file_bytes, filename)
    return parse_txt(file_bytes, filename)

